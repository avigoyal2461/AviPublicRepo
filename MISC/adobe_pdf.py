import pandas as pd
import os
from openpyxl import load_workbook
import openpyxl
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from PyPDF2 import PdfFileReader, PdfFileWriter
from reportlab.lib.pagesizes import A4, letter
#pip install pdfrw
import pdfrw
#pip install python-docx
import docx

# from PIL import Image, ImageOps

#NEEDS AN UPDATE, HOWEVER THE CODE TO WRITE INTO PDF FORM FIELDS IS READY, AT THE MOMENT ONLY INPUTS 2 POINTS OF DATA, HOWEVER THESE TWO VALUES ARE ALSO NOT THE SAME SELF VALUES
class write_field:

    def __init__(self):
        #cwd path, incase cwd does not work set manually..
        self.path = r"c:\adobe_test"
        #old paths
        # self.output = r"C:\adobe_test\adobe_test_final.pdf"
        # self.doc_output = r"C:\adobe_test\adobe_test_final.docx"
        # self.doc_output = os.path.join(self.path, "adobe_test_final.docx") # not needed...
        # self.loc = r"C:\adobe_test\Blank form data.xlsx"
        # self.pdf = r"C:\adobe_test\Blank_Form_template.pdf"
        # self.loc = r"C:\adobe_test\Adobe_test.xlsx"
        # self.pdf = r"C:\adobe_test\Adobe_test.pdf"

        #init paths
        self.pdf = os.path.join(self.path, "Blank_Form_template.pdf")
        self.output = os.path.join(self.path, "adobe_test_final.pdf")
        self.loc = os.path.join(self.path, "Sampledataset.xlsx")
        self.blank_pdf = os.path.join(self.path, "Blank.pdf")

        # self.xl = pd.ExcelFile(self.loc)
        self.xl = pd.read_excel(self.loc)
        self.set_columns()

        #create object to watch dimensons
        self.dimensions = self.xl.shape[0]
        self.names = []

        #excel rows
        self.Requested_Date = []
        self.Legal_Name = []
        self.Address = []
        self.Address_2 = []
        self.Email = []
        self.Phone = []
        self.First_Name = []
        self.Start_Date = []
        self.Reporting_Person_Name = []
        self.Days = []
        self.Submitted_By = []
        self.Department = []
        # print(self.dimensions)

    def set_columns(self) -> None:
        self.xl.columns = [
            'Requested_Date',
            'Legal_Name',
            'Address',
            'Address_2',
            'Email',
            'Phone',
            'First_Name',
            'Start_Date',
            'Reporting_Person_Name',
            'Days',
            'Submitted_By',
            'Department'
        ]
    def pull_all(self):
        """
        Runs all the pull requests
        """
        # one = self.pull_info_req()
        # two = self.pull_info_legalname()
        # three = self.pull_info_address()
        # four = self.pull_info_address2()
        # five = self.pull_info_email()
        # six = self.pull_info_phone()
        # seven = self.pull_info_firstname()
        # eight = self.pull_info_startdate()
        # nine = self.pull_info_reporting()
        # ten = self.pull_info_days()
        # eleven = self.pull_info_submitted()
        # twelve = self.pull_info_department()

        counter = 0
        for value in one:
            print(f"{one[counter]}, {two[counter]}, {three[counter]}, {four[counter]}, {five[counter]}, {six[counter]}, {seven[counter]}, {eight[counter]}, {nine[counter]}, {ten[counter]}, {eleven[counter]}, {twelve[counter]}")
            counter += 1

    def pull_info_req(self) -> list:
        """
        Pulls Requested Date
        """
        # for column in self.xl.columns:
        row_counter = 0
        Requested_Date = []
        temp = []
        while row_counter < self.dimensions:
            # print("something")
            date = self.xl.loc[row_counter].at['Requested_Date']
            temp.append(date)

            row_counter += 1

        for date in temp:
            date = str(date)
            temp_string = date.split(" ")
            self.Requested_Date.append(temp_string[0])
            Requested_Date.append(temp_string[0])
            # print(self.Requested_Date)

        return Requested_Date

    def pull_info_legalname(self) -> list:
        """
        Pulls Legal Name
        """
        # for column in self.xl.columns:
        row_counter = 0
        Legal_Name = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["Legal_Name"]
            self.Legal_Name.append(name)
            Legal_Name.append(name)
            row_counter += 1

        return Legal_Name

    def pull_info_address(self) -> list:
        """
        Pulls Address
        """
        # for column in self.xl.columns:
        row_counter = 0
        Address = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["Address"]
            self.Address.append(name)
            Address.append(name)
            row_counter += 1

        return Address

    def pull_info_address2(self) -> list:
        """
        Pulls Address_2
        """
        # for column in self.xl.columns:
        row_counter = 0
        Address_2 = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["Address_2"]
            self.Address_2.append(name)
            Address_2.append(name)
            row_counter += 1

        return Address_2

    def pull_info_email(self) -> list:
        """
        Pulls Email
        """
        # for column in self.xl.columns:
        row_counter = 0
        Email = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["Email"]
            self.Email.append(name)
            Email.append(name)
            row_counter += 1

        return Email

    def pull_info_phone(self) -> list:
        """
        Pulls Phone
        """
        # for column in self.xl.columns:
        row_counter = 0
        Phone = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["Phone"]
            self.Phone.append(name)
            Phone.append(name)
            row_counter += 1

        return Phone

    def pull_info_firstname(self) -> list:
        """
        Pulls First_Name
        """
        # for column in self.xl.columns:
        row_counter = 0
        First_Name = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["First_Name"]
            self.First_Name.append(name)
            First_Name.append(name)
            row_counter += 1

        return First_Name

    def pull_info_startdate(self) -> list:
        """
        Pulls Start_Date
        """
        # for column in self.xl.columns:
        row_counter = 0
        Start_Date = []
        temp = []
        while row_counter < self.dimensions:
            # print("something")
            date = self.xl.loc[row_counter].at['Start_Date']
            temp.append(date)

            row_counter += 1

        for date in temp:
            date = str(date)
            temp_string = date.split(" ")
            self.Start_Date.append(temp_string[0])
            Start_Date.append(temp_string[0])

        return Start_Date

    def pull_info_reporting(self) -> list:
        """
        Pulls Reporting_Person_Name
        """
        # for column in self.xl.columns:
        row_counter = 0
        Reporting_Person_Name = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["Reporting_Person_Name"]
            self.Reporting_Person_Name.append(name)
            Reporting_Person_Name.append(name)
            row_counter += 1

        return Reporting_Person_Name

    def pull_info_days(self) -> list:
        """
        Pulls Days
        """
        # for column in self.xl.columns:
        row_counter = 0
        Days = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["Days"]
            self.Days.append(name)
            Days.append(name)
            row_counter += 1

        return Days

    def pull_info_submitted(self) -> list:
        """
        Pulls Submitted_By
        """
        # for column in self.xl.columns:
        row_counter = 0
        Submitted_By = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["Submitted_By"]
            self.Submitted_By.append(name)
            Submitted_By.append(name)
            row_counter += 1

        return Submitted_By

    def pull_info_department(self) -> list:
        """
        Pulls Department
        """
        # for column in self.xl.columns:
        row_counter = 0
        Department = []
        while row_counter < self.dimensions:
            name = self.xl.loc[row_counter].at["Department"]
            self.Department.append(name)
            Department.append(name)
            row_counter += 1

        return Department

    def pull_info(self):
        """
        Pulls names into a list
        """
        row_counter = 0
        # names = []
        date = "temp"
        while date != "" or date != " " or date != "nan":
            date = self.df.loc[row_counter].at["Date"]
            # print(name)
            self.temp.append(date)
            for date in self.temp:
                date = str(date)
                temp = date.split(" ")
                self.date.append(temp[0])

            row_counter = row_counter + 1
            if row_counter == self.dimensions:
                break


        row_counter = 0
        legal_name = "temp"
        while legal_name != "" or legal_name != " " or legal_name != "nan":
            legal_name = self.df.loc[row_counter].at["Legal Name"]
            # print(name)
            self.legal_name.append(legal_name)
            row_counter = row_counter + 1
            if row_counter == self.dimensions:
                break

        row_counter = 0
        Approvername = "temp"
        while Approvername != "" or Approvername != " " or Approvername != "nan":
            Approvername = self.df.loc[row_counter].at["Approver name"]
            # print(name)
            self.approvername.append(Approvername)
            row_counter = row_counter + 1
            if row_counter == self.dimensions:
                break

        row_counter = 0
        submittedby = "temp"
        while submittedby != "" or submittedby != " " or submittedby != "nan":
            submittedby = self.df.loc[row_counter].at["Submitted by"]
            self.submittedby.append(submittedby)
            row_counter = row_counter + 1
            if row_counter == self.dimensions:
                break

        row_counter = 0
        phone = "temp"
        while phone != "" or phone != " " or phone != "nan":
            phone = self.df.loc[row_counter].at["Phone"]
            self.phone.append(phone)
            row_counter = row_counter + 1
            if row_counter == self.dimensions:
                break

        row_counter = 0
        email = "temp"
        while email != "" or email != " " or email != "nan":
            email = self.df.loc[row_counter].at["Email"]
            self.email.append(email)
            row_counter = row_counter + 1
            if row_counter == self.dimensions:
                break

        # print(self.email)
        return True
        #date, name, name of person ? , offer letter bs "position", for name of person, start date : blah blah

    def write_doc(self):
        # self.pull_info()
        # print(self.email)
        counter = 0
        for email in self.Email:
            # doc = aw.Document()

            # create a document builder object
            # builder = aw.DocumentBuilder(doc)
            builder = docx.Document()

            builder.add_paragraph(f"{self.Requested_Date[counter]}")
            # builder.add_paragraph("\n")
            builder.add_paragraph(f"{self.Legal_Name[counter]}")
            # builder.add_paragraph("\n")
            builder.add_paragraph(f"{self.Address[counter]}")

            if self.Address_2[counter] != "NaN" or self.Address_2[counter] != "nan":
                builder.add_paragraph(f"{self.Address_2[counter]}")

            # builder.add_paragraph("\n")
            builder.add_paragraph(f"{email}")
            # builder.add_paragraph(f"Submitted by: {self.submittedby[counter]}")
            # builder.add_paragraph("\n")
            builder.add_paragraph(f"{self.Phone[counter]}")
            builder.add_paragraph("\n")
            # builder.add_paragraph("\n")
            # builder.add_paragraph("\n")
            builder.add_paragraph(f"Dear {self.First_Name[counter]},")
            builder.add_paragraph("\n")
            # builder.add_paragraph("\n")
            # builder.add_paragraph("\n")
            builder.add_paragraph("We are pleased to offer you a position at Trinity Solar.")
            # builder.add_paragraph("\n")
            builder.add_paragraph(f"We would like you to start in that position from {self.Start_Date[counter]}.")
            builder.add_paragraph(f"You will be reporting to {self.Reporting_Person_Name[counter]} and will be working {self.Days[counter]}")
            builder.add_paragraph("\n")
            # builder.add_paragraph("\n")
            # builder.add_paragraph("\n")
            builder.add_paragraph("Thanks,")
            builder.add_paragraph("\n")
            # builder.add_paragraph("\n")
            # builder.add_paragraph("\n")
            builder.add_paragraph(f"Submitted by: {self.Submitted_By[counter]}")
            # builder.add_paragraph("\n")
            builder.add_paragraph(f"Department: {self.Department[counter]}")

            outname = f"{self.First_Name[counter]}_{str(self.Requested_Date[counter])}.docx"
            counter = counter + 1
            # output = os.path.join(self.path, outname, ".docx")

            builder.save(outname)

    def write_pdf_pixel(self):
        """
        IF THE PDF DOESNT HAVE FORM FIELDS
        """

        c = canvas.Canvas(self.blank_pdf, pagesize = letter)
        existing_pdf = PdfFileReader(open(self.pdf, "rb"))
        output = PdfFileWriter()

        page = existing_pdf.getPage(0)
        # output.addPage(page)

        # MAX_HEIGHT = 792
        MAX_HEIGHT = 600
        # MAX_WIDTH = 615
        #true max width is 615
        MAX_WIDTH = 500
        CENTER = MAX_WIDTH // 2
        MID = MAX_HEIGHT // 2
        QUARTER = MID // 2
        HALF_CENTER = CENTER // 2

        #set color to white, font type and size
        c.setFillColorRGB(255, 255, 255)
        c.setFont('Helvetica',20)

        # through_date = self.last_day_of_month()
        #starts at the left of page, written above the middle
        cen = 615 // 2
        c.drawCentredString(400, -300, "TEST DO I SEE THE TEXT ANYWHERRE HUI AUIGYUWEGFYUEFOUEAWOIAWEFUIOAUWAOYAWEFYUOWEFOYUAWFOUEWFUWOWEWOWEOQWOQUYG")
        # c.drawCentredString(cen, MID + 40, "TEST 2 " + "TEST 3" + str(self.date[1]))
        c.save()

        word_pdf_reader = PdfFileReader(self.blank_pdf, "rb")
        page.mergePage(word_pdf_reader.getPage(0))
        output.addPage(page)

        outputStream = open(self.output, 'wb')
        output.write(outputStream)
        outputStream.close()

        return True

    def write_pdf(self):

        self.pull_info()

        ANNOT_KEY = '/Annots'
        ANNOT_FIELD_KEY = '/T'
        ANNOT_VAL_KEY = '/V'
        ANNOT_RECT_KEY = '/Rect'
        SUBTYPE_KEY = '/Subtype'
        WIDGET_SUBTYPE_KEY = '/Widget'

        template_pdf = pdfrw.PdfReader(self.pdf)

        for page in template_pdf.pages:
            annotations = page[ANNOT_KEY]
            for annotation in annotations:
                if annotation[SUBTYPE_KEY] == WIDGET_SUBTYPE_KEY:
                    if annotation[ANNOT_FIELD_KEY]:
                        key = annotation[ANNOT_FIELD_KEY][1:-1]
                        # print(key)
                        # exit()

        for name in self.names:
            extension = "adobe_test_final_" + name + ".pdf"
            counter = 0
            email = self.email[counter]
            counter = counter + 1 #temp counter, iterates through the signature list without another loop.
            data_dict = {
        'Name': name,
        'Email': email
    }
            output = os.path.join(self.path, extension)
            self.fill_pdf(self.pdf, output, data_dict)

        return self.paths

    def fill_pdf(self, input_pdf_path, output_pdf_path, data_dict):

        ANNOT_KEY = '/Annots'
        ANNOT_FIELD_KEY = '/T'
        ANNOT_VAL_KEY = '/V'
        ANNOT_RECT_KEY = '/Rect'
        SUBTYPE_KEY = '/Subtype'
        WIDGET_SUBTYPE_KEY = '/Widget'

        template_pdf = pdfrw.PdfReader(input_pdf_path)
        for page in template_pdf.pages:
            annotations = page[ANNOT_KEY]
            for annotation in annotations:
                if annotation[SUBTYPE_KEY] == WIDGET_SUBTYPE_KEY:
                    if annotation[ANNOT_FIELD_KEY]:
                        key = annotation[ANNOT_FIELD_KEY][1:-1]
                        if key in data_dict.keys():
                            if type(data_dict[key]) == bool:
                                if data_dict[key] == True:
                                    annotation.update(pdfrw.PdfDict(
                                        AS=pdfrw.PdfName('Yes')))
                            else:
                                annotation.update(
                                    pdfrw.PdfDict(V='{}'.format(data_dict[key]))
                                )
                                annotation.update(pdfrw.PdfDict(AP=''))

        pdfrw.PdfWriter().write(output_pdf_path, template_pdf)
        # output_pdf_path = os.path.join(output_pdf_path, )
        self.paths.append(output_pdf_path)

if __name__ == "__main__":
    a = write_field()
    a.pull_all()
    a.write_doc()
    # a.run_doc()
    # a.pull_info_test()
    # a.write_pdf_pixel()
    # print(a.date)
    # print(a.categories)
    # print(a.write_pdf())
