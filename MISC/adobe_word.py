import pandas as pd
import os
from openpyxl import load_workbook
import openpyxl
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from PyPDF2 import PdfFileReader, PdfFileWriter
from reportlab.lib.pagesizes import A4, letter
#pip install pdfrw
import pdfrw
# import aspose.words as aw
#pip install aspose-words
#pip install python-docx
import docx

# from PIL import Image, ImageOps
#rewrite where it says signature to email
class write_field:

    def __init__(self):
        self.path = r"c:\adobe_test"
        # self.path = r"C:\Users\RPA_Bot_7\Desktop\adobe_test"
        # self.output = r"C:\adobe_test\adobe_test_final.pdf"
        # self.output = r"C:\Users\RPA_Bot_7\Desktop\adobe_test\adobe_test_final.pdf"
        # self.doc_output = r"C:\adobe_test\adobe_test_final.docx"
        # self.doc_output = r"C:\Users\RPA_Bot_7\Desktop\adobe_test\adobe_test_final.docx"
        self.loc = r"C:\adobe_test\Adobe_test.xlsx"
        # self.loc = r"C:\Users\RPA_Bot_7\Desktop\adobe_test\Adobe_test.xlsx"
        # self.pdf = r"C:\adobe_test\Adobe_test.pdf"
        # self.pdf = r"C:\Users\RPA_Bot_7\Desktop\adobe_test\Adobe_test.pdf"
        self.xl = pd.ExcelFile(self.loc)
        self.df = self.xl.parse()
        self.dimensions = self.df.shape
        self.dimensions = self.dimensions[0]
        self.names = []
        self.signatures = []
        self.paths = []
        # print(self.dimensions)

    def pull_info(self):
        """
        Pulls names into a list
        """
        row_counter = 0
        # names = []
        name = "temp"
        while name != "" or name != " " or name != "nan":
            name = self.df.loc[row_counter].at["Full Name"]
            print(name)
            self.names.append(name)
            row_counter = row_counter + 1
            if row_counter == self.dimensions:
                break

        row_counter = 0

        signature = "temp"
        while signature != "" or signature != " " or signature != "nan":
            signature = self.df.loc[row_counter].at["Email"]
            self.signatures.append(signature)
            row_counter = row_counter + 1
            if row_counter == self.dimensions:
                break

        return True

    def build_doc(self, Name, First, Last):
        """
        builds a doc with python docx
        """
        output = f"{self.path}/{First}{Last}.docx"
        # output = os.path.join(self.path, output, ".docx")
        doc = docx.Document()
        print("Writing...")
        doc.add_paragraph(f"Hello {Name}")
        doc.add_paragraph("\n")
        doc.add_paragraph("Please sign here...")
        doc.save(output)


                    # document.merge([{'Requested_Date' : self.adobe.Requested_Date[counter]},
                    #                {'Legal_Name' : self.adobe.Legal_Name[counter]},
                    #                {'Address' : self.adobe.Address[counter]},
                    #                {'Address_2' : self.adobe.Address_2[counter]},
                    #                {'Email' : email},
                    #                {'Phone' : self.adobe.Phone[counter]},
                    #                {'First_name' : self.adobe.First_Name[counter]},
                    #                {'Start_date' : self.adobe.Start_Date[counter]},
                    #                {'Reporting_Person_Name' : self.adobe.Reporting_Person_Name[counter]},
                    #                {'Days' : self.adobe.Days[counter]},
                    #                {'Submitted_By' : self.adobe.Submitted_By[counter]},
                    #                {'Department' : self.adobe.Department[counter]}],)
                    # document.merge(document_dict)
                    # document.merge(Name=self.adobe.First_Name[counter], Total = "{:,.2f}".format())

    # def build_doc(self, Name, First, Last):
    #     """
    #     builds a doc with python docx
    #     """
    #     output = f"{self.path}/{First}{Last}.docx"
    #     # output = os.path.join(self.path, output, ".docx")
    #     doc = docx.Document()
    #     print("Writing...")
    #     doc.add_paragraph(f"Hello {Name}")
    #     doc.add_paragraph("\n")
    #     doc.add_paragraph("Please sign here...")
    #     doc.save(output)
    #     self.paths.append(output)
    #
    #
    # def run_doc(self):
    #     self.pull_info()
    #
    #     for name in self.names:
    #         list_name = name.split(" ")
    #         first = list_name[0]
    #         last = list_name[1]
    #         # self.write_doc_aspose(first, last)
    #         self.build_doc(name, first, last)
    #         # print(first)
    #         # print(last)

if __name__ == "__main__":
    a = write_field()
    a.run_doc()
    # print(a.write_pdf())

    # def write_doc(self, First, Last):
    #     # self.pull_info()
    #
    #     doc = aw.Document()
    #
    #     # create a document builder object
    #     builder = aw.DocumentBuilder(doc)
    #
    #     # create font
    #     font = builder.font
    #     font.size = 12
    #     # font.bold = True
    #     font.name = "Arial"
    #     # font.underline = aw.Underline.DASH
    #
    #     # set paragraph formatting
    #     paragraphFormat = builder.paragraph_format
    #     paragraphFormat.first_line_indent = 8
    #     paragraphFormat.alignment = aw.ParagraphAlignment.JUSTIFY
    #     paragraphFormat.keep_together = True
    #
    #     # add text
    #     builder.writeln(f"Hello {First} {Last}")
    #     builder.writeln("\n")
    #     builder.writeln("Please sign here...")
    #
    #     # save document
    #     outname = f"{First}{Last}"
    #     outname = f"{First}{Last}.docx"
    #
    #     output = os.path.join(self.path, outname, ".docx")
    #
    #     doc.save(outname)
